"""
Import July 2026 duty roster from DOCX into ShiftRoster.

Source file:
  C:/Users/Administrator/Desktop/JULY 2026 DUTY ROSTER.docx

Run:
  python import_july_roster.py

Shift leaders (fixed by management):
  A -> Ssewanyana Ronald
  B -> Suuna Emmanuel
  C -> Musambi Isaac
  D -> Mugisha William
"""

import os
from datetime import datetime
from pathlib import Path

import docx

os.chdir(os.path.dirname(os.path.abspath(__file__)))

from app import create_app, db
from app.models.apron import ShiftRoster
from app.models.user import User


DOCX_PATH = Path(r"C:/Users/Administrator/Desktop/JULY 2026 DUTY ROSTER.docx")

# Canonical shift membership list with corrected names.
SHIFTS = {
    "A": [
        "Ssewanyana Ronald",
        "Byaruhanga Edgar",
        "Ssemwanga Ronald",
        "Alele Abel",
        "Opolot Emmanuel",
        "Namugenyi Robinah",
    ],
    "B": [
        "Suuna Emmanuel",
        "Musinguzi Victor",
        "Kyakonye Ziad",
        "Sharif Mohamed",
        "Ederu Edgar",
        "Ecuru Albert",
        "Abiti Mary Florence",
    ],
    "C": [
        "Musambi Isaac",
        "Ogola John",
        "Ssebuwufu Francis",
        "Akandwanaho Stella",
        "Mukisa Edward",
        "Bakashaba Jackson",
    ],
    "D": [
        "Mugisha William",
        "Candiga Francis",
        "Asewu Moses",
        "Ssentongo Ivan",
        "Matovu Mark",
        "Evon Akandwanaho",
    ],
}

SHIFT_LEADERS = {
    "A": "Ssewanyana Ronald",
    "B": "Suuna Emmanuel",
    "C": "Musambi Isaac",
    "D": "Mugisha William",
}

DUTY_MAP = {
    "D": "day",
    "N": "night",
    "O": "off",
}

CYCLE_INDEX = {
    "day": 0,
    "night": 1,
    "off": 2,
}


def _normalize(name: str) -> str:
    return " ".join((name or "").strip().lower().split())


def _find_or_create_user(full_name: str):
    norm = _normalize(full_name)

    user = User.query.filter(
        db.func.lower(db.func.trim(User.full_name)) == norm
    ).first()
    if user:
        return user, False

    parts = [p for p in norm.split() if len(p) >= 2]
    query = User.query
    for part in parts:
        query = query.filter(User.full_name.ilike(f"%{part}%"))
    user = query.first()
    if user:
        return user, False

    base_username = "".join(ch for ch in norm.split()[0] if ch.isalnum()) or "user"
    username = base_username
    idx = 1
    while User.query.filter_by(username=username).first():
        username = f"{base_username}{idx}"
        idx += 1

    email = f"{username}@airside.ebb"
    eidx = 1
    while User.query.filter_by(email=email).first():
        email = f"{username}{eidx}@airside.ebb"
        eidx += 1

    user = User(
        username=username,
        email=email,
        full_name=full_name,
        role="operator",
        department="Airside Operations",
        is_active=True,
    )
    user.set_password("Airside@2026!")
    db.session.add(user)
    db.session.flush()
    return user, True


def _parse_docx_daily_schedule(path: Path):
    if not path.exists():
        raise FileNotFoundError(f"Roster file not found: {path}")

    doc = docx.Document(str(path))
    if not doc.tables:
        raise RuntimeError("No table found in July roster DOCX.")

    table = doc.tables[0]
    rows = []

    # Data rows begin after table headers.
    for r in range(6, len(table.rows)):
        values = [table.cell(r, c).text.strip().replace("\n", " ") for c in range(len(table.columns))]
        raw_date = values[1].strip() if len(values) > 1 else ""
        if not raw_date:
            continue

        # Example values: 01-07-26, 31-7-26
        parsed = None
        for fmt in ("%d-%m-%y", "%d-%m-%Y"):
            try:
                parsed = datetime.strptime(raw_date, fmt).date()
                break
            except ValueError:
                continue

        if parsed is None:
            # Fallback parser for single-digit month/day formatting.
            parts = raw_date.split("-")
            if len(parts) == 3:
                day = int(parts[0])
                month = int(parts[1])
                year = int(parts[2])
                if year < 100:
                    year += 2000
                parsed = datetime(year, month, day).date()
            else:
                raise ValueError(f"Could not parse roster date: {raw_date}")

        a_code = values[2].strip().upper()
        b_code = values[3].strip().upper()
        c_code = values[4].strip().upper()
        d_code = values[5].strip().upper()

        if a_code in DUTY_MAP and b_code in DUTY_MAP and c_code in DUTY_MAP and d_code in DUTY_MAP:
            rows.append((parsed, a_code, b_code, c_code, d_code))

    if len(rows) != 31:
        raise RuntimeError(f"Expected 31 schedule rows for July, got {len(rows)}")

    return rows


def run():
    app = create_app("development")
    with app.app_context():
        admin = User.query.filter_by(role="admin").first()
        admin_id = admin.id if admin else None

        print("\n=== Parsing July roster document ===")
        daily_schedule = _parse_docx_daily_schedule(DOCX_PATH)
        print(f"  Parsed {len(daily_schedule)} day rows from {DOCX_PATH}")

        print("\n=== Resolving shift members to user accounts ===")
        shift_users = {}
        created_users = []
        for shift_letter, members in SHIFTS.items():
            shift_users[shift_letter] = []
            for name in members:
                user, was_created = _find_or_create_user(name)
                shift_users[shift_letter].append(user)
                if was_created:
                    created_users.append(user)
                state = "CREATED" if was_created else "matched"
                leader_tag = " [LEADER]" if _normalize(name) == _normalize(SHIFT_LEADERS[shift_letter]) else ""
                print(f"  Shift {shift_letter}: {state:<7} -> {user.full_name}{leader_tag}")

        db.session.commit()

        print("\n=== Upserting ShiftRoster entries (July 2026) ===")
        created_count = 0
        updated_count = 0
        skipped_count = 0

        shift_order = ["A", "B", "C", "D"]
        for item in daily_schedule:
            duty_date = item[0]
            duty_codes = {shift_order[i]: item[i + 1] for i in range(4)}

            for shift_letter, code in duty_codes.items():
                duty_type = DUTY_MAP[code]
                cycle_idx = CYCLE_INDEX[duty_type]
                leader_name = SHIFT_LEADERS[shift_letter]

                for user in shift_users[shift_letter]:
                    existing = ShiftRoster.query.filter_by(user_id=user.id, duty_date=duty_date).first()
                    if existing:
                        if existing.duty_type in ("leave", "study_leave", "office"):
                            skipped_count += 1
                            continue
                        existing.duty_type = duty_type
                        existing.cycle_day_index = cycle_idx
                        existing.notes = f"july2026_roster|shift_{shift_letter}|leader={leader_name}"
                        existing.created_by_user_id = admin_id
                        updated_count += 1
                    else:
                        db.session.add(
                            ShiftRoster(
                                duty_date=duty_date,
                                user_id=user.id,
                                duty_type=duty_type,
                                cycle_day_index=cycle_idx,
                                notes=f"july2026_roster|shift_{shift_letter}|leader={leader_name}",
                                created_by_user_id=admin_id,
                            )
                        )
                        created_count += 1

        db.session.commit()

        print("\n=== July roster import complete ===")
        print(f"  Roster entries created : {created_count}")
        print(f"  Roster entries updated : {updated_count}")
        print(f"  Skipped (leave/office) : {skipped_count}")

        if created_users:
            print(f"\n  New user accounts created ({len(created_users)}):")
            for user in created_users:
                print(f"    - {user.full_name} ({user.username})")
            print("  Temporary password for new users: Airside@2026!")


if __name__ == "__main__":
    run()
